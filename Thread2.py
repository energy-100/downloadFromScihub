import os
import sys
from functools import reduce
from PyQt5.QtGui import *
from PyQt5.QtCore import *
from PyQt5 import QtCore
import traceback
import re
import pickle
import os
import datetime
import requests
import urllib
import json
import socket
import PyPDF2
import PyPDF4
from docx import Document
import urllib3
from googletrans import Translator
from bs4 import BeautifulSoup

from google_trans_new import google_translator

from multiprocessing.dummy import Pool as ThreadPool

os.environ['REQUESTS_CA_BUNDLE'] =  os.path.join(os.path.dirname(sys.argv[0]), 'cacert.pem')
# from requests.packages.urllib3.exceptions import InsecureRequestWarning
from resumable import urlretrieve
import urllib3
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
# requests.packages.urllib3.disable_warnings(InsecureRequestWarning)
# urllib3.disable_warnings()
# requests.packages.urllib3.disable_warnings()
import ssl
ssl._create_default_https_context = ssl._create_unverified_context


import importlib, sys

importlib.reload(sys)
from pdfminer.pdfparser import PDFParser, PDFDocument
from pdfminer.pdfdevice import PDFDevice
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import PDFPageAggregator
from pdfminer.layout import LTTextBoxHorizontal, LAParams
from pdfminer.pdfinterp import PDFTextExtractionNotAllowed
# from pdfminer.pdfpage import PDFPage

class artobject():
    def __init__(self, parent=None):
        self.filename = ""
        self.path = ""
        self.allpath = ""
        self.data = ""
        self.datastr = ""

        self.title = ""
        self.indextime = ""
        self.deposittime = ""
        self.updatetime = ""
        self.createdtime = ""
        self.publishedtime = ""
        self.online_published = ""
        self.print_published = ""
        self.posted = ""
        self.accepted = ""
        self.type=""
        self.source=""
        self.reference_count=-1
        self.is_reference_count=-1
        self.source=""
        self.titlerecode = ""
        self.chinesetitle = ""
        self.extra = ""
        self.doi = ""
        self.authors = ""
        self.arturl = ""
        self.infrequesturl = ""
        self.abstract = ""
        self.abstracttrans = ""
        self.extractpath=""
        self.transpath=""

class updateurlsThread(QThread):
    messageSingle = QtCore.pyqtSignal(str)
    enddingSingle = QtCore.pyqtSignal(list)
    def __init__(self,allurl):
        self.allurl=allurl
        self.ableurls=[]
        super(updateurlsThread, self).__init__()
    def run(self):
        '''请求头'''
        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/65.0.3325.181 Safari/537.36",
        }

        '''代理ip'''
        proxies = {
            "http": "http://202.121.96.33:8086",
            "http": "http://58.253.154.179:9999",
            "http": "http://163.204.243.51:9999",
            "http": "http://183.166.20.179:9999",
            "http": "http://183.166.20.179:9999",
            "http": "http://49.86.181.35:9999",

            # "https": "https://221.228.17.172:8181",

        }

        urls = [
            "https://sci-hub.tw",
            "https://sci-hub.se",
            "https://sci-hub.st",
            "https://sci-hub.pl",
            "https://sci-hub.ren",
            "https://sci-hub.sci-hub.hk",
            "https://sci-hub.hk",
            "https://sci-hub.la",
            "https://sci-hub.mn",
            "https://sci-hub.name",
            "https://sci-hub.is",
            "https://sci-hub.tv",
            "https://sci-hub.ws",
            "https://www.sci-hub.cn",
            "https://sci-hub.sci-hub.tw",
            "https://sci-hub.sci-hub.mn",
            "https://sci-hub.sci-hub.tv"]
        print("正在更新线路...")
        ableurls = []

        for i in range(len(urls)):
            try:
                currentPag_html = requests.get(urls[i], headers=headers, timeout=10,verify=False)  # 访问该网站
            except Exception as a:
                print(urls[i] + "失败！("+str(a)+")")
                self.messageSingle.emit(urls[i].split(".")[-1] + "（"+str(i)+"/"+str(len(urls))+"）线路失败！")
                continue
            # 先判断是否成功访问
            if currentPag_html.status_code == 200:
                print(urls[i] + "成功！")
                # self.messageSingle.emit(urls[i].split(".")[-1] + "线路成功！")
                # delay=ping_host(urls[i])
                # print(urls[i]+"成功！ 延迟：",delay)
                self.ableurls.append(urls[i])
            else:
                print(urls[i] + "失败！")
                # self.messageSingle.emit(urls[i] + "失败！")
            if len(self.ableurls)>=1:
                break
        self.enddingSingle.emit(self.ableurls)

class updateurlsThreadsingle(QThread):
    messageSingle = QtCore.pyqtSignal(str)
    enddingSingle = QtCore.pyqtSignal(str,float)
    def __init__(self,url,timeout):
        self.time=9999
        self.url=url
        self.timeout=timeout
        self.ableurls=[]
        super(updateurlsThreadsingle, self).__init__()
    def run(self):
        '''请求头'''
        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/65.0.3325.181 Safari/537.36",
        }

        '''代理ip'''
        proxies = {
            "http": "http://202.121.96.33:8086",
            "http": "http://58.253.154.179:9999",
            "http": "http://163.204.243.51:9999",
            "http": "http://183.166.20.179:9999",
            "http": "http://183.166.20.179:9999",
            "http": "http://49.86.181.35:9999",

            # "https": "https://221.228.17.172:8181",

        }



        print("正在更新线路...")
        try:
            if self.url in ["https://www.crossref.org/","https://translate.google.cn/"]:
                currentPag_html = requests.get(self.url, headers=headers, timeout=self.timeout,verify=False)  # 访问该网站
            else:
                testurldata={"request": "Quantum mechanics and hidden variables: A test of Bell's inequality by the measurement of the spin correlation in low-energy proton-proton scattering"}
                currentPag_html = requests.post(self.url, data=testurldata, verify=False,headers=headers, timeout=self.timeout)
                main_html = currentPag_html.text
                soup = BeautifulSoup(main_html, 'html.parser')
                pdfurl = soup.find_all('iframe')
                if len(pdfurl)==0:
                    currentPag_html.text=""
        except Exception as a:
            # print(self.url + "失败！("+str(a)+")")
            # self.messageSingle.emit(self.url.split(".")[-1] +"）线路失败！  "+str(a))
            self.enddingSingle.emit(self.url, self.time)
            return
        # 先判断是否成功访问
        if currentPag_html.status_code == 200:
            if currentPag_html.text=="":
                self.time=10000
                print(self.url + "返回值为空！" + str(currentPag_html.status_code))
                # self.messageSingle.emit(self.url + "失败！" + str(currentPag_html.status_code))
            else:
                # print(self.url + "成功！")
                self.time=currentPag_html.elapsed.total_seconds()
                # self.messageSingle.emit(self.url.split(".")[-1] + "线路成功！("+str(self.time)+")")
                # delay=ping_host(urls[i])
                # print(urls[i]+"成功！ 延迟：",delay)

        else:
            print(self.url + "失败！"+str(currentPag_html.status_code))
            # self.messageSingle.emit(self.url + "失败！"+str(currentPag_html.status_code))
        # if len(self.ableurls)>=1:
        #     return
        self.enddingSingle.emit(self.url,self.time)

class extractThread(QThread):
    messageSingle = QtCore.pyqtSignal(str)
    enddingSingle = QtCore.pyqtSignal()
    progressSingle = QtCore.pyqtSignal(int)
    progressvisualSingle = QtCore.pyqtSignal(bool)
    def __init__(self,art:artobject):
        self.art=art
        super(extractThread, self).__init__()

    def valid_xml_char_ordinal(self,c):
        codepoint = ord(c)
        # conditions ordered by presumed frequency
        return (
                0x20 <= codepoint <= 0xD7FF or
                codepoint in (0x9, 0xA, 0xD) or
                0xE000 <= codepoint <= 0xFFFD or
                0x10000 <= codepoint <= 0x10FFFF
        )

    def run(self):
        self.progressvisualSingle.emit(True)
        self.progressSingle.emit(0)
        self.messageSingle.emit("正在进行OCR识别...")
        pdfFile = open(self.art.allpath, 'rb')
        # pdfReader = PyPDF4.PdfFileReader(pdfFile)
        # print(pdfReader.getDocumentInfo())
        # print(pdfReader.numPages)
        content=""
        # for i in range(pdfReader.numPages):
        #     content+=pdfReader.getPage(i).extractText()
        #     self.progressSingle.emit(int((i+1)/pdfReader.numPages*100))
        # content = content.replace('\r', ' ').replace('\n', ' ').replace('\t', ' ')
        # pdfFile.close()


        parser = PDFParser(pdfFile)
        # 创建一个PDF文档
        doc = PDFDocument()
        # 分析器和文档相互连接
        parser.set_document(doc)
        doc.set_parser(parser)
        # 提供初始化密码，没有默认为空
        doc.initialize()
        # 检查文档是否可以转成TXT，如果不可以就忽略
        if not doc.is_extractable:
            raise PDFTextExtractionNotAllowed
        else:
            # 创建PDF资源管理器，来管理共享资源
            rsrcmagr = PDFResourceManager()
            # 创建一个PDF设备对象
            laparams = LAParams()
            # 将资源管理器和设备对象聚合
            device = PDFPageAggregator(rsrcmagr, laparams=laparams)
            # 创建一个PDF解释器对象
            interpreter = PDFPageInterpreter(rsrcmagr, device)

            # 循环遍历列表，每次处理一个page内容
            # doc.get_pages()获取page列表
            pagenum=sum(1 for _ in doc.get_pages())
            for i,page in enumerate(doc.get_pages()):
                # print(len(doc.get_pages()))
                interpreter.process_page(page)
                # 接收该页面的LTPage对象
                layout = device.get_result()
                # 这里的layout是一个LTPage对象 里面存放着page解析出来的各种对象
                # 一般包括LTTextBox，LTFigure，LTImage，LTTextBoxHorizontal等等一些对像
                # 想要获取文本就得获取对象的text属性
                for x in layout:
                    try:
                        if (isinstance(x, LTTextBoxHorizontal)):
                            # with open('%s' % (save_path), 'a') as f:
                            result = x.get_text()
                            # print(result)
                            content += result
                                # f.write(result + "\n")
                    except:
                        print("Failed")
                self.progressSingle.emit(int((i + 1) / pagenum * 100))
            content = content.replace('\r', ' ').replace('\t', ' ').replace('.\n', '*---*---*').replace('\n',' ').replace('*---*---*', u'\r\n\r\n')

        self.messageSingle.emit("正在保存...")

        document = Document()                          # 打开一个基于默认“模板”的空白文档
        document.add_heading(self.art.title, 0)      # 添加标题
        # contentdoc1=content.encode("utf-8")
        # contentdoc=str(content.encode("utf-8"))
        # p = document.add_paragraph(str(content.encode("utf-8")))
        content = ''.join(c for c in content if self.valid_xml_char_ordinal(c))
        p = document.add_paragraph(content)
        try:
            path=self.art.path + "/" + self.art.titlerecode + ".docx"
            document.save(path)
        except PermissionError :
            self.progressvisualSingle.emit(False)
            self.messageSingle.emit("文件占用，请关闭后重试！")
        except Exception as e:
            self.messageSingle.emit("未知异常，请重试！("+str(e)+")")
        self.art.extractpath=path
        self.progressvisualSingle.emit(False)
        self.messageSingle.emit("文本提取完成！")
        self.enddingSingle.emit()

class transThread(QThread):
    messageSingle = QtCore.pyqtSignal(str)
    enddingSingle = QtCore.pyqtSignal()
    progressSingle = QtCore.pyqtSignal(int)
    progressvisualSingle = QtCore.pyqtSignal(bool)
    def __init__(self,art:artobject,translator):
        self.art=art
        self.translator=translator
        # self.pool = ThreadPool(8)
        super(transThread, self).__init__()
    def run(self):
        # self.progressvisualSingle.emit(True)
        # self.progressSingle.emit(0)
        self.messageSingle.emit("正在翻译...")
        self.progressvisualSingle.emit(True)
        self.progressSingle.emit(0)
        document = Document(self.art.extractpath)
        content=""
        for paragraph in document.paragraphs:
            content += paragraph.text
            # print(len(paragraph.text))
        # content=content.replace('\r',' ').replace('\n',' ').replace('\t',' ')
        sentencelist=re.split(r"([.])", content)
        chineseword=""
        paragraphlist=[]
        tempparagraph=""
        for i,sentence in enumerate(sentencelist):
            tempparagraph+=sentence
            if len(tempparagraph)>1000:
                # translator = Translator()
                try:
                    # chineseword += self.translator.translate(tempparagraph[0:(len(tempparagraph)-len(sentence))], 'zh-CN').text
                    # tmpstr=self.pool.map(self.request, str(tempparagraph[0:(len(tempparagraph)-len(sentence))]))
                    # chineseword = chineseword.join(tmpstr)
                    chineseword += self.translator.translate(tempparagraph[0:(len(tempparagraph)-len(sentence))], 'zh-CN')
                except Exception as e:
                    print(e)
                    self.messageSingle.emit("翻译线路超时，请检查网络连接或稍后重试！(" + str(e) + ")")
                tempparagraph=sentence
            self.progressSingle.emit((i+1)/len(sentencelist)*100)
        try:
            chineseword += self.translator.translate(tempparagraph, 'zh-CN')
            # chineseword += self.translator.translate(tempparagraph, 'zh-CN').text
        except Exception as e:
            print(e)
            self.messageSingle.emit("翻译线路超时，请检查网络连接或稍后重试！(" + str(e) + ")")
        # self.pool.close()
        # self.pool.join()
        documentchs = Document()                          # 打开一个基于默认“模板”的空白文档
        # documentchs.add_heading(self.art.title, 0)      # 添加标题
        p = documentchs.add_paragraph(chineseword)
        try:
            path=self.art.path + "/" + self.art.titlerecode + "_chs.docx"
            documentchs.save(path)
            self.art.transpath = path
        except PermissionError :
            # self.progressvisualSingle.emit(False)
            self.messageSingle.emit("文件占用，请关闭后重试！")
        except Exception as e:
            self.messageSingle.emit("未知异常，请重试！("+str(e)+")")
        finally:
            self.progressvisualSingle.emit(False)
        # self.progressvisualSingle.emit(False)
        self.messageSingle.emit("文本翻译完成！")
        self.enddingSingle.emit()

    # def request(self,text):
    #     lang = 'zh-CN'
    #     # t = google_translator(timeout=5)
    #     translate_text = self.translator.translate(text.strip(), lang)
    #     return translate_text


class getartThread(QThread):
    messageSingle = QtCore.pyqtSignal(str)
    enddingSingle = QtCore.pyqtSignal(list)
    def __init__(self,urls,title,elemunmperpage,trans,timetype,starttime,endtime,translator,page=1):
        self.timetype=timetype
        self.starttime=starttime
        self.endtime=endtime
        self.elemunmperpage=elemunmperpage+1
        self.urls = urls
        self.title = title
        self.translator=translator
        self.page = page
        self.trans=trans
        super(getartThread, self).__init__()

    def run(self):
        headers = {"User-Agent": "GroovyBib/1.1 (mailto:abc551218@126.com) BasedOnFunkyLib/1.4",'Connection':'close'}

        '''代理ip'''
        proxies = \
            {
            "http": "http://202.121.96.33:8086",
            "http": "http://58.253.154.179:9999",
            "http": "http://163.204.243.51:9999",
            "http": "http://183.166.20.179:9999",
            "http": "http://183.166.20.179:9999",
            "http": "http://49.86.181.35:9999",
            # "https": "https://221.228.17.172:8181",
            }
        startdatafilterlist=[
        "",
        "from-index-date",
        "from-deposit-date",
        "from-update-date",
        "from-created-date",
        "from-pub-date",
        "from-online-pub-date",
        "from-print-date",
        "from-posted-date",
        "from-accepted-date",
        ]

        enddatafilterlist=[
        "",
        ",until-index-date",
        ",until-deposit-date",
        ",until-update-date",
        ",until-created-date",
        ",until-pub-date",
        ",until-online-pub-date",
        ",until-print-date",
        ",until-posted-date",
        ",until-accepted-date",
        ]

        self.messageSingle.emit("正在请求数据...")
        filterstr=""
        if self.timetype!=0:
            filterstr+="&filter=is-update:true,"+startdatafilterlist[self.timetype]+":"+self.starttime+""+enddatafilterlist[self.timetype]+":"+self.endtime

        try:
            # title = "Quantum mechanics and hidden variables: A test of Bell's inequality by the measurement of the spin correlation in low-energy proton-proton scattering"
            url = "https://api.crossref.org/works?query=" + self.title +"&rows="+str(self.elemunmperpage)+"&offset="+str((self.page-1)*self.elemunmperpage)+filterstr
            rawdata = requests.get(url, headers=headers, timeout=100,verify=False)  # 访问该网站
        except Exception as a:
            self.messageSingle.emit("请求数据错误（"+str(a)+")")
            return

        try:
            self.messageSingle.emit("正在解析数据...")
            print("rawdata.text:",rawdata.text)
            data = json.loads(rawdata.text)
            dataitems = data.get("message").get("items")
            # translator = Translator()
            artlist=[]
            for i,artitem in enumerate(dataitems):
                self.messageSingle.emit("正在获取数据 " + str(i + 1) + "/" + str(len(dataitems)))
                # title=artitem.find("p",class_="lead").string
                # print(title)
                artobjecttemp = artobject()
                if artitem.get("title"):
                    artobjecttemp.title=str(re.sub("\n","",artitem.get("title")[0]))
                else :
                    artobjecttemp.title=artitem.get("DOI")

                if self.trans:
                    # print(artobjecttemp.title)
                    # translator = Translator()
                    try:
                        artobjecttemp.chinesetitle = self.translator.translate(artobjecttemp.title, 'zh-CN')
                        # artobjecttemp.chinesetitle = self.translator.translate(artobjecttemp.title, dest='zh-CN').text
                    except Exception as e:
                        print(e)
                        artobjecttemp.chinesetitle="翻译线路超时"
                else:
                    artobjecttemp.chinesetitle="翻译未开启！"
                # artobjecttemp.extra = artitem.find("p", class_="extra").get_text().strip().replace("\n", " ")
                if artitem.get("author"):
                    for elem in artitem.get("author"):
                        if elem.get("given"):
                            artobjecttemp.authors += elem["given"] + "-" + elem["family"] + "(" + elem["sequence"] + "),"
                        elif elem.get("name"):
                            artobjecttemp.authors += elem["name"] + "(" + elem["sequence"] + "),"
                else:
                    artobjecttemp.authors="Not Found!"
                if artitem.get("source"):
                    artobjecttemp.source = artitem.get("source")

                if artitem.get("created"):
                    artobjecttemp.createdtime=artitem.get("created").get("date-time")
                if artitem.get("indexed"):
                    artobjecttemp.indextime=artitem.get("indexed").get("date-time")
                if artitem.get("deposited").get("date-time"):
                    artobjecttemp.deposittime = artitem.get("deposited").get("date-time")
                if artitem.get("update"):
                    artobjecttemp.updatetime = artitem.get("update").get("date-parts")[0][0] if len(artitem.get("update").get("date-parts")[0])==1 else reduce(lambda x, y: str(x) + "-" + str(y),artitem.get("update").get("date-parts")[0])
                if artitem.get("issued"):
                    artobjecttemp.publishedtime = artitem.get("issued").get("date-parts")[0][0] if len(artitem.get("issued").get("date-parts")[0])==1 else reduce(lambda x, y: str(x) + "-" + str(y),artitem.get("issued").get("date-parts")[0])
                if artitem.get("published-online"):
                    artobjecttemp.online_published = artitem.get("published-online").get("date-parts")[0][0] if len(artitem.get("published-online").get("date-parts")[0])==1 else reduce(lambda x, y: str(x) + "-" + str(y),artitem.get("published-online").get("date-parts")[0])
                if artitem.get("published-print"):
                    artobjecttemp.print_published = artitem.get("published-print").get("date-parts")[0][0] if len(artitem.get("published-print").get("date-parts")[0])==1 else reduce(lambda x, y: str(x) + "-" + str(y),artitem.get("published-print").get("date-parts")[0])
                if artitem.get("posted"):
                    artobjecttemp.posted = artitem.get("posted").get("date-parts")[0][0] if len(artitem.get("posted").get("date-parts")[0])==1 else reduce(lambda x, y: str(x) + "-" + str(y),artitem.get("posted").get("date-parts")[0])
                if artitem.get("accepted"):
                    artobjecttemp.accepted = artitem.get("accepted").get("date-parts")[0][0] if len(artitem.get("accepted").get("date-parts")[0])==1 else reduce(lambda x, y: str(x) + "-" + str(y),artitem.get("accepted").get("date-parts")[0])




                container_title=""
                volume=""
                issue=""
                page=""
                if artitem.get("container-title"):
                    container_title=artitem.get("container-title")[0]
                if artitem.get("volume"):
                    volume=" volume:"+artitem.get("volume")
                if artitem.get("issue"):
                    issue=" issue:"+artitem.get("issue")
                if artitem.get("page"):
                    page=" page:"+artitem.get("page")
                if artitem.get("abstract"):
                    artobjecttemp.abstract = artitem.get("abstract")
                    # translator = Translator()
                    try:
                        artobjecttemp.abstracttrans += self.translator.translate(artobjecttemp.abstract,'zh-CN')
                        # artobjecttemp.abstracttrans += self.translator.translate(artobjecttemp.abstract,'zh-CN').text
                    except Exception as e:
                        print(e)
                        self.messageSingle.emit("翻译线路超时，请检查网络连接或稍后重试！(" + str(e) + ")")

                artobjecttemp.extra="published in "+container_title+volume+issue+page
                # if artitem.get("volume"):

                if artitem.get("is-referenced-by-count"):
                    artobjecttemp.is_reference_count=artitem.get("is-referenced-by-count")

                if artitem.get("references-count"):
                    artobjecttemp.reference_count = artitem.get("references-count")

                if artitem.get("type"):
                    artobjecttemp.type = artitem.get("type")

                artobjecttemp.arturl=artitem.get("URL")
                # if arturl is not None:
                #     artobjecttemp.arturl = arturl

                # if artobjecttemp.arturl !="":
                artobjecttemp.doi = artitem.get("DOI")

                # if artobjecttemp.doi !="":
                #     artobjecttemp.infrequesturl= "https://api.crossref.org/works/" + artobjecttemp.doi

                artlist.append(artobjecttemp)
        except Exception as a:
            self.messageSingle.emit("解析数据错误（"+str(traceback.format_exc(limit=1))+")")
            print(traceback.format_exc(limit=1))
            # self.messageSingle.emit("解析数据错误（"+str(a)+")")
            return
        self.enddingSingle.emit(artlist)

        # try:
        #     # title = "Quantum mechanics and hidden variables: A test of Bell's inequality by the measurement of the spin correlation in low-energy proton-proton scattering"
        #     url = "https://search.crossref.org/?q=" + self.title +"&page="+str(self.page)
        #     currentPag_html = requests.get(url, headers=headers, timeout=100,verify=False)  # 访问该网站
        # except Exception as a:
        #     self.messageSingle.emit("请求数据错误（"+str(a)+")")
        #     return
        # try:
        #     self.messageSingle.emit("正在解析数据...")
        #     main_html = currentPag_html.text
        #     soup = BeautifulSoup(main_html, 'html.parser')
        #     artitems = soup.find_all('td', class_="item-data")
        #     artlist=[]
        #     translator = Translator()
        #     for i,artitem in enumerate(artitems):
        #         self.messageSingle.emit("正在获取数据 " + str(i + 1) + "/" + str(len(artitems)))
        #         title=artitem.find("p",class_="lead").string
        #         # print(title)
        #         artobjecttemp = artobject()
        #         lead=artitem.find("p", class_="lead")
        #         title=lead.find("title")
        #         if title:
        #             artobjecttemp.title = title.string.strip()
        #         else:
        #             artobjecttemp.title = lead.string.strip()
        #         if self.trans:
        #             artobjecttemp.chinesetitle = translator.translate(artobjecttemp.title, dest='zh-CN').text
        #         else:
        #             artobjecttemp.chinesetitle="翻译线路超时！"
        #         artobjecttemp.extra = artitem.find("p", class_="extra").get_text().strip().replace("\n", " ")
        #         authors = artitem.find("p", class_="expand")
        #         if authors is not None:
        #             artobjecttemp.authors = authors.get_text().strip().replace("\n", " ")[8:]
        #
        #         arturl=artitem.find('a')['href']
        #         if arturl is not None:
        #             artobjecttemp.arturl = arturl
        #
        #         if artobjecttemp.arturl !="":
        #             artobjecttemp.doi = artobjecttemp.arturl[16:]
        #
        #         if artobjecttemp.doi !="":
        #             artobjecttemp.infrequesturl= "https://api.crossref.org/works/" + artobjecttemp.doi
        #
        #         artlist.append(artobjecttemp)
        # except Exception as a:
        #     self.messageSingle.emit("解析数据错误（"+str(i)+str(traceback.format_exc(limit=1))+")")
        #     print(str(i)+traceback.format_exc(limit=1))
        #     # self.messageSingle.emit("解析数据错误（"+str(a)+")")
        #     return
        # self.enddingSingle.emit(artlist)

        # arturl = soup.find_all('div', class_="item-links")
        # authors = soup.find_all('p', class_="expand")
        # authors = authors[0].string
        # print(authors)
        # arturl = arturl[0].find('a')['href']
        # doi = arturl[16:]
        # artinfurl = "https://api.crossref.org/works/" + doi
        # artinfjson = requests.get(artinfurl, headers=headers, timeout=20)  # 访问该网站
        # artinfdict = json.loads(artinfjson.text)
        # print(artinfdict["message"])
        # for key, value in artinfdict["message"].items():
        #     print(key, value)
        # scihuburl = "https://sci-hub.tw/" + doi
        # scihub_html = requests.get(scihuburl, headers=headers, timeout=20)  # 访问该网站
        # main_html = scihub_html.text
        # soup = BeautifulSoup(main_html, 'html.parser')
        # pdfurl = soup.find('iframe')['src']
        # urllib.request.urlretrieve(pdfurl, 'C:/Users/ENERGY/Desktop/download/555.pdf')
        # print(pdfurl)
        # self.enddingSingle.emit(self.savepath,self.title + ".pdf")

class savefilethread(QThread):
    messageSingle = QtCore.pyqtSignal(str)
    enddingSingle = QtCore.pyqtSignal(artobject)
    codeimageSingle = QtCore.pyqtSignal(QImage)
    progressSingle = QtCore.pyqtSignal(int)
    progressvisualSingle = QtCore.pyqtSignal(bool)

    def __init__(self,scihuburl:list,art:artobject,path:str):
        self.scihuburl = scihuburl
        self.art = art
        self.path = path
        self.erroormessage=""
        super(savefilethread, self).__init__()


    def Schedule(self,a,b,c):
        #a:已经下载的数据块
        # b:数据块的大小
        #c:远程文件的大小
        per = 100.0 * a * b / c
        if per > 100 :
            per = 100
        # mess='已完成：'+ str(round(per,2)+"%%")
        # self.messageSingle.emit(mess)
        self.progressSingle.emit(int(per))
    def run(self):
        headers = {
            'User-Agent': 'Mozilla/5.0 (X11; Linux x86_64; rv:27.0) Gecko/20100101 Firefox/27.0',
        }

        '''代理ip'''
        proxies = \
            {
                "http": "http://202.121.96.33:8086",
                "http": "http://58.253.154.179:9999",
                "http": "http://163.204.243.51:9999",
                "http": "http://183.166.20.179:9999",
                "http": "http://183.166.20.179:9999",
                "http": "http://49.86.181.35:9999",
                # "https": "https://221.228.17.172:8181",
            }
        filepath=""
        for i,url in enumerate(self.scihuburl):
            try:
                self.messageSingle.emit("正在使用"+url+"("+str(i+1)+"/"+str(len(self.scihuburl))+")建立连接")
                # scihuburl = self.scihuburl+"/" + self.art.doi
                # scihub_html = requests.get(scihuburl, headers=headers, timeout=20)  # 访问该网站
                # scihub_html = requests.post(self.scihuburl,{"request":self.art.arturl},headers=headers, verify = False)
                if self.art.doi==self.art.title:
                    scihub_html = requests.get(url+"/"+self.art.doi, headers=headers, verify=False)
                else:
                    scihub_html = requests.post(url,{"request":self.art.doi},headers=headers, verify = False)
                # scihub_html = requests.post(self.scihuburl,{"request":self.art.title},headers=headers, verify = False)
                # scihub_html = requests.post(self.scihuburl,{"request":self.art.title})
                print(scihub_html.status_code)
                main_html = scihub_html.text
                soup = BeautifulSoup(main_html, 'html.parser')
                # identifyingcode=soup.find_all(id='captcha')
                # if len(identifyingcode)>0:
                #     self.erroormessage="下载失败！"
                #     print("验证码")
                #     continue
                #     imageurl=identifyingcode[0]['src']
                #     # file = cStringIO.StringIO(urllib2.urlopen(self.scihuburl+imageurl).read())
                #     # img = Image.open(file)
                #     res = requests.get(self.scihuburl+imageurl,verify=False)
                #     img = QImage.fromData(res.content)
                #
                #     self.codeimageSingle.emit(img)
                #
                # # pdfurl = soup.find('iframe')['src']

                pdfurl = soup.find_all('iframe')
                if len(pdfurl)==0:
                    print("self.scihuburl",self.scihuburl)
                    self.messageSingle.emit("sci-hub未收录此文章，无法下载！")
                    self.erroormessage="sci-hub未收录此文章，无法下载！"
                    continue
                rstr = r"[\/\\\:\*\?\"\<\>\|]"
                filename=re.sub(rstr, "_", self.art.title)
                if not os.path.exists(self.path):
                    os.makedirs(self.path)
                filepath=self.path+"/"+filename+".pdf"
                self.messageSingle.emit("正在下载...（由于网络问题可能会消耗较长时间）")
                if pdfurl[0]['src'][0:4]!="http":
                    pdfurl2="https:"+pdfurl[0]['src']
                else:
                    pdfurl2 = pdfurl[0]['src']
                print(pdfurl2)
                self.progressSingle.emit(0)
                self.progressvisualSingle.emit(True)
                socket.setdefaulttimeout(100)
                urlretrieve(pdfurl2, filepath,self.Schedule)
                # urllib.request.urlretrieve(pdfurl2, filepath,self.Schedule)
                self.progressvisualSingle.emit(False)
                self.art.filename = self.art.title+".pdf"
                self.art.titlerecode = filename
                self.art.path = self.path
                self.art.allpath = filepath
                self.art.data = datetime.datetime.now()
                self.art.datastr = datetime.datetime.now().strftime('%y-%m-%d %H:%M:%S')
                self.messageSingle.emit("文章下载成功！")
                self.erroormessage="文章下载成功！"
                self.enddingSingle.emit(self.art)
                break
            except PermissionError:
                # print(e)
                print("self.scihuburl", url)
                print("文章下载失败！（文件占用，请关闭pdf文件后尝试下载)")  # 或者加入其它判断条件
                self.messageSingle.emit("文章下载失败！（文件占用，请关闭pdf文件后尝试下载)")
                self.erroormessage="文章下载失败！（文件占用，请关闭pdf文件后尝试下载)"
                break
            except Exception as a :
                print(repr(a))
                print("self.scihuburl", url)
                self.messageSingle.emit("文章下载失败,请稍后重试！（"+str(a)+")")
                self.erroormessage="文章下载失败,请稍后重试！（"+str(a)+")"
                self.progressvisualSingle.emit(False)
                if os.path.exists(filepath):
                    try:
                        os.remove(filepath)
                    except Exception as a:
                        pass
                continue
        self.progressvisualSingle.emit(False)
        self.messageSingle.emit(self.erroormessage)

class saveachethread(QThread):
    messageSingle = QtCore.pyqtSignal(str)
    enddingSingle = QtCore.pyqtSignal(list)
    def __init__(self,achepath,arts):
        self.achepath = achepath
        self.arts=arts
        super(saveachethread, self).__init__()
    def run(self):
        achefilepath, achefilename = os.path.split(self.achepath)
        if (not os.path.exists(achefilepath)):
            os.makedirs(achefilepath)
        with open(self.achepath, "wb") as file:
            pickle.dump(self.arts, file, True)

class translatethread(QThread):
    messageSingle = QtCore.pyqtSignal(str)
    enddingSingle = QtCore.pyqtSignal(str)
    def __init__(self,enstr,translator):
        self.enstr = enstr
        self.translator=translator
        super(translatethread, self).__init__()
    def run(self):
        chineseword=""
        try:
            chineseword=self.translator.translate(self.enstr,'zh-CN')
            # chineseword=self.translator.translate(self.enstr,'zh-CN').text
        except Exception as e:
            print(e)
            chineseword="翻译超时！"
        self.enddingSingle.emit(chineseword)

class getauthoritythread(QThread):
    messageSingle = QtCore.pyqtSignal(str)
    enddingSingle = QtCore.pyqtSignal(str)
    def __init__(self,enstr,translator):
        self.enstr = enstr
        self.translator=translator
        super(getauthoritythread, self).__init__()
    def run(self):
        # translator = Translator()
        try:
            # chineseword=self.translator.translate(self.enstr, 'zh-CN').text
            chineseword=self.translator.translate(self.enstr, 'zh-CN')
        except Exception as e:
            print(e)
            chineseword=""
        self.enddingSingle.emit(chineseword)

